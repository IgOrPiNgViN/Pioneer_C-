# -*- coding: utf-8 -*-
"""Генерация контрольной работы (темы 1-17) в Word — два варианта."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


def add_answer_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run("Ответ: ").bold = True
    p.add_run("_" * 72)


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(code.strip("\n"))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def build_variant1() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Контрольная работа по C++\n"
        "Темы уроков 1–17 (основы + ООП)\n"
        "Вариант 1"
    )
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "Формат: 10 вопросов — 5 теоретических (часть А), 5 практических — с кодом (часть Б). "
        "В части А дайте краткие ответы: число, слово, «да/нет» или одно короткое предложение "
        "на каждый подпункт. В части Б — вывод программы и/или исправление ошибки в 1–3 предложениях."
    )

    doc.add_heading("Часть А. Теоретические вопросы (краткие ответы)", level=1)

    theory = [
        (
            "Вопрос А1 (уроки 1–2: введение, переменные и типы).",
            "1) С какой функции начинается выполнение программы на C++ (имя)?\n"
            "2) Литерал 7 без точки — это значение типа ...?\n"
            "3) Оператор вывода в консоль через std::cout — два символа: ...?",
        ),
        (
            "Вопрос А2 (уроки 3–4: операторы, условия).",
            "1) Чему равно значение выражения 2 + 3 * 4 (одно число)?\n"
            "2) Какой оператор используют для проверки на равенство двух чисел?\n"
            "3) Ключевое слово «иначе» в разветвлении if — на английском: ...?",
        ),
        (
            "Вопрос А3 (уроки 5–7: циклы, массивы, матрицы).",
            "1) Сколько раз выполнится тело цикла while (false) { ... } (число)?\n"
            "2) У массива int a[8] последний допустимый индекс — ...?\n"
            "3) В записи b[i][j] первый индекс обычно обозначает ... (строку или столбец)?",
        ),
        (
            "Вопрос А4 (уроки 8–12: строки, функции, указатели, структуры, продвинутые функции).",
            "1) Как у std::string узнать длину (имя метода)?\n"
            "2) Несколько функций с одним именем, но разными параметрами — это ... (одно слово-термин)?\n"
            "3) В объявлении void f(int& x) параметр x передаётся по ... (значению или ссылке)?",
        ),
        (
            "Вопрос А5 (уроки 13–17: файлы, классы, конструкторы, наследование, полиморфизм).",
            "1) Класс из <fstream> для чтения из файла: std::...?\n"
            "2) Сокрытие полей класса и доступ к ним только через методы — принцип ООП ...?\n"
            "3) Чтобы при вызове через указатель Base* выполнялась версия метода из производного класса, "
            "в базовом методе пишут ключевое слово ...?",
        ),
    ]

    for head, body in theory:
        doc.add_paragraph(head, style="List Number")
        doc.add_paragraph(body)
        add_answer_line(doc)

    doc.add_heading("Часть Б. Практические вопросы (код)", level=1)

    practice = [
        (
            "Вопрос Б1 (уроки 2–3). Что выведет фрагмент? Объясните порядок вычисления.",
            """#include <iostream>
using namespace std;
int main() {
    int a = 10, b = 4;
    cout << a / b << " " << a % b << endl;
    cout << (a > b && b++ == 4) << endl;
    cout << b << endl;
    return 0;
}""",
        ),
        (
            "Вопрос Б2 (уроки 4–5). Найдите логическую ошибку: программа должна печатать "
            "числа 1…5, но поведение неверное. Укажите строку и исправление.",
            """#include <iostream>
using namespace std;
int main() {
    int i = 1;
    while (i < 5) {
        std::cout << i << " ";
        i = i + 1;
    }
    return 0;
}""",
        ),
        (
            "Вопрос Б3 (уроки 6–8). Что выведет программа?",
            """#include <iostream>
#include <string>
using namespace std;
int main() {
    string s = "ab";
    s += s;
    cout << s.size() << " " << s << endl;
    return 0;
}""",
        ),
        (
            "Вопрос Б4 (уроки 9–11). Найдите ошибку компиляции или времени выполнения и исправьте.",
            """#include <iostream>
using namespace std;
int twice(int x) { return 2 * x; }

int main() {
    int p = nullptr;
    cout << twice(*p) << endl;
    return 0;
}""",
        ),
        (
            "Вопрос Б5 (уроки 14–17). Найдите ошибку в определении класса (компиляция) и исправьте.",
            """#include <iostream>
using namespace std;

class Point {
    int x, y
public:
    Point(int a, int b) : x(a), y(b) {}
    int getX() const { return x; }
};

int main() {
    Point p(1, 2);
    cout << p.getX() << endl;
    return 0;
}""",
        ),
    ]

    for head, code in practice:
        doc.add_paragraph(head, style="List Number")
        add_code_block(doc, code)

    return doc


def build_variant2() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Контрольная работа по C++\n"
        "Темы уроков 1–17 (основы + ООП)\n"
        "Вариант 2"
    )
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "Формат: 10 вопросов — 5 теоретических (часть А), 5 практических с кодом (часть Б). "
        "В части А — только краткие ответы. В части Б — вывод и/или исправление ошибки кратко."
    )

    doc.add_heading("Часть А. Теоретические вопросы (краткие ответы)", level=1)

    theory = [
        (
            "Вопрос А1 (уроки 1–2).",
            "1) Программа на C++ перед запуском переводится компилятором в ... код (одно слово: "
            "исходный / машинный / графический)?\n"
            "2) Заголовок для std::cin и std::cout: #include <...>?\n"
            "3) Можно ли безопасно выводить int k; до первого присваивания k (да или нет)?",
        ),
        (
            "Вопрос А2 (уроки 3–4).",
            "1) Результат целочисленного выражения 17 / 5 для типа int — ...?\n"
            "2) Остаток от деления 17 % 5 — ...?\n"
            "3) Оператор выбора по значению целой переменной — ключевое слово ...?",
        ),
        (
            "Вопрос А3 (уроки 5–7).",
            "1) В цикле do { ... } while (условие); тело выполняется минимум ... раз (число)?\n"
            "2) У массива int m[5][6] второй размер (6) — это число ... (строк или столбцов)?\n"
            "3) Ключевое слово немедленного выхода из цикла — ...?",
        ),
        (
            "Вопрос А4 (уроки 8–12).",
            "1) Тип «стрелка на int» в записи int* p — это ... (указатель / ссылка / массив)?\n"
            "2) Ключевое слово для объявления структуры в C++ — ...?\n"
            "3) Локальная переменная внутри функции видна ... (где: везде в файле / только в этой функции)?",
        ),
        (
            "Вопрос А5 (уроки 13–17).",
            "1) Метод is_open() у файлового потока проверяет, ... ли файл (одно слово: открыт / закрыт / удалён)?\n"
            "2) Специальная функция класса с именем как у класса, создающая объект — ...?\n"
            "3) Запись class B : public A означает, что B — ... класса A (одно слово)?",
        ),
    ]

    for head, body in theory:
        doc.add_paragraph(head, style="List Number")
        doc.add_paragraph(body)
        add_answer_line(doc)

    doc.add_heading("Часть Б. Практические вопросы (код)", level=1)

    practice = [
        (
            "Вопрос Б1 (уроки 2–4). Что выведет программа?",
            """#include <iostream>
using namespace std;
int main() {
    int n = 2;
    if (n = 3)
        cout << "A";
    else
        cout << "B";
    cout << " " << n << endl;
    return 0;
}""",
        ),
        (
            "Вопрос Б2 (уроки 5–6). Найдите ошибку: программа должна суммировать элементы "
            "массива из 3 чисел {2, 3, 4}, но результат неверный.",
            """#include <iostream>
using namespace std;
int main() {
    int a[3] = {2, 3, 4};
    int sum = 0;
    for (int i = 0; i <= 3; ++i)
        sum += a[i];
    cout << sum << endl;
    return 0;
}""",
        ),
        (
            "Вопрос Б3 (уроки 7–9). Что выведет программа?",
            """#include <iostream>
using namespace std;

int f(int x) { return x + 1; }

int main() {
    int a[2][2] = {{1, 2}, {3, 4}};
    cout << a[1][0] << f(0) << endl;
    return 0;
}""",
        ),
        (
            "Вопрос Б4 (уроки 10–12). Найдите ошибку компиляции и исправьте.",
            """#include <iostream>
using namespace std;

void swap(int a, int b) {
    int t = a; a = b; b = t;
}

int main() {
    int x = 1, y = 9;
    swap(x, y);
    cout << x << " " << y << endl;
    return 0;
}""",
        ),
        (
            "Вопрос Б5 (уроки 14–17). Программа не компилируется. Укажите минимальное исправление.",
            """#include <iostream>
using namespace std;

class Animal {
public:
    virtual void speak() { cout << "..." << endl; }
};

class Dog : public Animal {
public:
    void speak() const override { cout << "woof" << endl; }
};

int main() {
    Dog d;
    Animal* p = &d;
    p->speak();
    return 0;
}""",
        ),
    ]

    for head, code in practice:
        doc.add_paragraph(head, style="List Number")
        add_code_block(doc, code)

    return doc


def _save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        alt = path.with_name(path.stem + "_копия.docx")
        doc.save(alt)
        return alt


def main() -> None:
    base = Path(__file__).resolve().parent
    v1 = base / "контрольная_темы_1-17_вариант_1.docx"
    v2 = base / "контрольная_темы_1-17_вариант_2.docx"
    p1 = _save_doc(build_variant1(), v1)
    p2 = _save_doc(build_variant2(), v2)
    print("Saved:", p1)
    print("Saved:", p2)


if __name__ == "__main__":
    main()
